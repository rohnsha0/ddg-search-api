from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import List, Optional, Dict
from io import BytesIO
import os

# Professional color scheme - Light and subtle
COLOR_PRIMARY = RGBColor(70, 100, 130)  # Soft Blue-Gray
COLOR_SECONDARY = RGBColor(100, 125, 150)  # Muted Blue-Gray
COLOR_ACCENT = RGBColor(245, 246, 248)  # Very Light Neutral
COLOR_TEXT = RGBColor(80, 80, 80)  # Soft Gray
COLOR_LIGHT_BG = RGBColor(252, 252, 252)  # Almost White
COLOR_HEADER_BG = RGBColor(240, 242, 245)  # Light Gray-Blue for headers
COLOR_HEADER_TEXT = RGBColor(70, 100, 130)  # Soft Blue-Gray text
COLOR_BORDER = RGBColor(230, 232, 235)  # Very Light Gray


class WeeklyStatusReportGenerator:
    """Class to generate professional weekly status reports in DOCX format"""
    
    def __init__(self):
        """Initialize the report generator"""
        pass
    
    def add_border_to_table(self, table):
        """Add minimal professional borders to table (horizontal lines only)"""
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        # Add table borders
        tblBorders = OxmlElement('w:tblBorders')
        
        # Only horizontal borders with light color
        for border_name in ['top', 'bottom', 'insideH']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '6')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'E6E8EB') # Very light gray
            tblBorders.append(border)
            
        # Remove vertical borders
        for border_name in ['left', 'right', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')
            tblBorders.append(border)
            
        tblPr.append(tblBorders)
        self.keep_table_rows_together(table)
    

    def keep_table_rows_together(self, table):
        """Prevent table rows from breaking across pages"""
        for row in table.rows:
            tr = row._element
            trPr = tr.get_or_add_trPr()
            # Keep row together on one page
            cantSplit = OxmlElement('w:cantSplit')
            trPr.append(cantSplit)

    def create_generic_table(self, doc, section_title, headers, data_rows, column_widths=None):
        """
        Create a professionally formatted table with given headers and data
        
        Args:
            doc: Document object
            section_title: Title for the section
            headers: List of column headers
            data_rows: List of dictionaries with data for each row
            column_widths: List of width values in Inches (optional)
        """
        # Section heading
        section_heading = doc.add_heading(section_title, level=2)
        section_heading.paragraph_format.space_before = Pt(10)
        section_heading.paragraph_format.space_after = Pt(8)
        section_heading_run = section_heading.runs[0]
        section_heading_run.font.color.rgb = COLOR_PRIMARY
        section_heading_run.font.size = Pt(12)
        section_heading_run.font.name = 'Calibri'
        
        # Create table
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        header_cells = table.rows[0].cells
        
        # Format header row
        for i, header in enumerate(headers):
            cell = header_cells[i]
            cell.text = header
            
            # Set cell padding - minimal
            tcPr = cell._element.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for margin_name in ['top', 'left', 'bottom', 'right']:
                margin = OxmlElement(f'w:{margin_name}')
                margin.set(qn('w:w'), '50')
                margin.set(qn('w:type'), 'dxa')
                tcMar.append(margin)
            tcPr.append(tcMar)
            
            # Format header
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.text = header
            run = paragraph.runs[0]
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = COLOR_HEADER_TEXT
            run.font.name = 'Calibri'
            
            # Header background color - light
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F0F2F5')
            cell._element.get_or_add_tcPr().append(shading_elm)
        
        # Set column widths
        if column_widths:
            for i, width in enumerate(column_widths):
                for row in table.rows:
                    row.cells[i].width = width
        
        # Add data rows
        for item in data_rows:
            row_cells = table.add_row().cells
            
            # Populate cells with data
            for i, header in enumerate(headers):
                key = header.lower().replace(' ', '_').replace('/', '_').replace('(', '').replace(')', '')
                row_cells[i].text = str(item.get(key, ''))
            
            # Format data cells
            for i, cell in enumerate(row_cells):
                # Add cell padding - minimal
                tcPr = cell._element.get_or_add_tcPr()
                tcMar = OxmlElement('w:tcMar')
                for margin_name in ['top', 'left', 'bottom', 'right']:
                    margin = OxmlElement(f'w:{margin_name}')
                    margin.set(qn('w:w'), '50')
                    margin.set(qn('w:type'), 'dxa')
                    tcMar.append(margin)
                tcPr.append(tcMar)
                
                paragraph = cell.paragraphs[0]
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)
                
                # Center alignment for specific columns
                if 'complete' in headers[i].lower() or 'impact' in headers[i].lower():
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                paragraph.text = cell.text
                run = paragraph.runs[0]
                run.font.size = Pt(8)
                run.font.color.rgb = COLOR_TEXT
                run.font.name = 'Calibri'
        
        self.add_border_to_table(table)
        doc.add_paragraph()  # Spacing after table
    
    def generate(
        self,
        project_name: str,
        report_date: str,
        progress_data: List[Dict[str, str]],
        client_logo_path: Optional[str] = None,
        company_logo_path: Optional[str] = None,
        tasks_completed: Optional[List[Dict[str, str]]] = None,
        planned_next_week: Optional[List[Dict[str, str]]] = None,
        risks_issues: Optional[List[Dict[str, str]]] = None,
        blockers: Optional[List[Dict[str, str]]] = None,
        milestone_status: Optional[List[Dict[str, str]]] = None,
        decisions_needed: Optional[List[Dict[str, str]]] = None,
        dependencies: Optional[List[Dict[str, str]]] = None,
        action_items: Optional[List[Dict[str, str]]] = None
    ) -> BytesIO:
        """
        Generate a weekly status report and return as BytesIO object
        
        Args:
            project_name: Name of the project
            report_date: Date string (e.g., "Week Ending: December 11, 2024")
            progress_data: List of dictionaries with keys: 'area', 'planned', 'completed', 'percent', 'notes'
            client_logo_path: Path to client logo image (optional)
            company_logo_path: Path to company logo image (optional)
            tasks_completed: List of dictionaries with keys: 'tasks', 'owner', 'notes'
            planned_next_week: List of dictionaries with keys: 'task', 'owner', 'expected_outcome'
            risks_issues: List of dictionaries with keys: 'risk_issue', 'impact_hml', 'owner', 'mitigation'
            blockers: List of dictionaries with keys: 'blockers', 'owner', 'action_reqd'
            milestone_status: List of dictionaries with keys: 'milestone', 'planned_date', 'status', 'comment'
            decisions_needed: List of dictionaries with keys: 'decision', 'impact', 'due_by'
            dependencies: List of dictionaries with keys: 'dependency', 'status', 'owner'
            action_items: List of dictionaries with keys: 'action_item', 'owner', 'due_by', 'status'
        
        Returns:
            BytesIO object containing the DOCX document
        """
        doc = Document()
        
        # Set default logo paths if they exist
        if not client_logo_path and os.path.exists("client_logo.png"):
            client_logo_path = "client_logo.png"
        if not company_logo_path and os.path.exists("company_logo.png"):
            company_logo_path = "company_logo.png"
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(1.1)
            section.right_margin = Inches(1.1)
        
        # === HEADER WITH LOGOS ===
        if client_logo_path or company_logo_path:
            header_para = doc.add_paragraph()
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_para.paragraph_format.space_before = Pt(6)
            header_para.paragraph_format.space_after = Pt(12)
            
            # Create a table with no borders for logo positioning
            header_table = doc.add_table(rows=1, cols=3)
            header_table.autofit = False
            
            # Remove all borders from header table
            tbl = header_table._element
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            
            # Set table to have no borders
            tblBorders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                tblBorders.append(border)
            tblPr.append(tblBorders)
            
            # Client logo (left)
            if client_logo_path:
                cell_left = header_table.rows[0].cells[0]
                cell_left.width = Inches(1.8)
                paragraph_left = cell_left.paragraphs[0]
                paragraph_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
                shading_left = OxmlElement('w:shd')
                shading_left.set(qn('w:fill'), 'FFFFFF')
                cell_left._element.get_or_add_tcPr().append(shading_left)
                try:
                    run_left = paragraph_left.add_run()
                    run_left.add_picture(client_logo_path, width=Inches(1.3))
                except Exception:
                    pass
            
            # Center cell - empty for spacing
            cell_center = header_table.rows[0].cells[1]
            cell_center.width = Inches(2.4)
            shading_center = OxmlElement('w:shd')
            shading_center.set(qn('w:fill'), 'FFFFFF')
            cell_center._element.get_or_add_tcPr().append(shading_center)
            
            # Company logo (right)
            if company_logo_path:
                cell_right = header_table.rows[0].cells[2]
                cell_right.width = Inches(1.8)
                paragraph_right = cell_right.paragraphs[0]
                paragraph_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                shading_right = OxmlElement('w:shd')
                shading_right.set(qn('w:fill'), 'FFFFFF')
                cell_right._element.get_or_add_tcPr().append(shading_right)
                try:
                    run_right = paragraph_right.add_run()
                    run_right.add_picture(company_logo_path, width=Inches(1.3))
                except Exception:
                    pass
            
            doc.add_paragraph()  # Spacing
        
        # === REPORT TITLE ===
        title = doc.add_heading('WEEKLY STATUS REPORT', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(18)
        title_run.font.color.rgb = COLOR_PRIMARY
        title_run.font.bold = True
        title_run.font.name = 'Calibri'
        
        # === PROJECT INFO ===
        project_para = doc.add_paragraph()
        project_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        project_para.paragraph_format.space_before = Pt(4)
        project_para.paragraph_format.space_after = Pt(10)
        project_run = project_para.add_run(f'{project_name}\n{report_date}')
        project_run.font.size = Pt(11)
        project_run.font.color.rgb = COLOR_TEXT
        project_run.font.name = 'Calibri'
        project_run.bold = True
        
        doc.add_paragraph()  # Spacing
        
        # === PROGRESS VS PLAN SECTION ===
        section_heading = doc.add_heading('Progress vs Plan', level=2)
        section_heading.paragraph_format.space_before = Pt(10)
        section_heading.paragraph_format.space_after = Pt(8)
        section_heading_run = section_heading.runs[0]
        section_heading_run.font.color.rgb = COLOR_PRIMARY
        section_heading_run.font.size = Pt(12)
        section_heading_run.font.name = 'Calibri'
        
        # === CREATE PROGRESS TABLE ===
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        header_cells = table.rows[0].cells
        headers = ['Area/Workstream', 'Planned This Week', 'Completed', '% Complete', 'Notes']
        
        for i, header in enumerate(headers):
            cell = header_cells[i]
            cell.text = header
            
            # Set cell padding - minimal
            tcPr = cell._element.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for margin_name in ['top', 'left', 'bottom', 'right']:
                margin = OxmlElement(f'w:{margin_name}')
                margin.set(qn('w:w'), '50')
                margin.set(qn('w:type'), 'dxa')
                tcMar.append(margin)
            tcPr.append(tcMar)
            
            # Format header
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(3)
            
            # Replace runs by setting paragraph text to avoid direct XML manipulation
            paragraph.text = header
            run = paragraph.runs[0]
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = COLOR_HEADER_TEXT
            run.font.name = 'Calibri'
            
            # Header background color - light
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F0F2F5')
            cell._element.get_or_add_tcPr().append(shading_elm)
        
        # Set column widths
        widths = [Inches(1.9), Inches(1.9), Inches(1.4), Inches(0.9), Inches(1.9)]
        for i, width in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = width
        
        # Add data rows
        for item in progress_data:
            row_cells = table.add_row().cells
            row_cells[0].text = item.get('area', '')
            row_cells[1].text = item.get('planned', '')
            row_cells[2].text = item.get('completed', '')
            row_cells[3].text = item.get('percent', '')
            row_cells[4].text = item.get('notes', '')
            
            # Format data cells
            for i, cell in enumerate(row_cells):
                # Add cell padding - minimal
                tcPr = cell._element.get_or_add_tcPr()
                tcMar = OxmlElement('w:tcMar')
                for margin_name in ['top', 'left', 'bottom', 'right']:
                    margin = OxmlElement(f'w:{margin_name}')
                    margin.set(qn('w:w'), '50')
                    margin.set(qn('w:type'), 'dxa')
                    tcMar.append(margin)
                tcPr.append(tcMar)
                
                paragraph = cell.paragraphs[0]
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)
                
                # Alignment by column
                if i == 3:  # Percent column
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Set paragraph text directly (avoids XML manipulation) then style its run
                paragraph.text = cell.text
                run = paragraph.runs[0]
                run.font.size = Pt(8)
                run.font.color.rgb = COLOR_TEXT
                run.font.name = 'Calibri'
        
        self.add_border_to_table(table)
        doc.add_paragraph()  # Spacing after table
        
        # === TASKS COMPLETED SECTION ===
        if tasks_completed:
            self.create_generic_table(
                doc,
                'Tasks Completed',
                ['Tasks', 'Owner', 'Notes'],
                tasks_completed,
                column_widths=[Inches(2.5), Inches(1.5), Inches(2.5)]
            )
        
        # === PLANNED FOR NEXT WEEK SECTION ===
        if planned_next_week:
            self.create_generic_table(
                doc,
                'Planned for Next Week',
                ['Task', 'Owner', 'Expected Outcome'],
                planned_next_week,
                column_widths=[Inches(2.5), Inches(1.5), Inches(2.5)]
            )
        
        # === RISKS / ISSUES & MITIGATIONS SECTION ===
        if risks_issues:
            self.create_generic_table(
                doc,
                'Risks / Issues & Mitigations',
                ['Risk/Issue', 'Impact (H/M/L)', 'Owner', 'Mitigation'],
                risks_issues,
                column_widths=[Inches(2.0), Inches(1.2), Inches(1.3), Inches(2.0)]
            )
        
        # === BLOCKERS SECTION ===
        if blockers:
            self.create_generic_table(
                doc,
                'Blockers',
                ['Blockers', 'Owner', 'Action Reqd'],
                blockers,
                column_widths=[Inches(2.5), Inches(1.5), Inches(2.5)]
            )
        
        # === MILESTONE STATUS SECTION ===
        if milestone_status:
            self.create_generic_table(
                doc,
                'Milestone Status',
                ['Milestone', 'Planned Date', 'Status', 'Comment'],
                milestone_status,
                column_widths=[Inches(2.0), Inches(1.4), Inches(1.2), Inches(2.4)]
            )
        
        # === DECISIONS NEEDED FROM CLIENT SECTION ===
        if decisions_needed:
            self.create_generic_table(
                doc,
                'Decisions Needed from Client',
                ['Decision', 'Impact', 'Due By'],
                decisions_needed,
                column_widths=[Inches(2.8), Inches(2.2), Inches(2.0)]
            )
        
        # === DEPENDENCIES SECTION ===
        if dependencies:
            self.create_generic_table(
                doc,
                'Dependencies',
                ['Dependency', 'Status', 'Owner'],
                dependencies,
                column_widths=[Inches(2.8), Inches(1.8), Inches(2.4)]
            )
        
        # === ACTION ITEMS SECTION ===
        if action_items:
            self.create_generic_table(
                doc,
                'Action Items',
                ['Action Item', 'Owner', 'Due By', 'Status'],
                action_items,
                column_widths=[Inches(2.2), Inches(1.5), Inches(1.3), Inches(1.5)]
            )
        
        # === FOOTER ===
        doc.add_paragraph()
        
        # Footer content
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.paragraph_format.space_after = Pt(0)
        footer_run = footer.add_run('Confidential - Internal Use Only')
        footer_run.font.size = Pt(8)
        footer_run.font.italic = True
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        footer_run.font.name = 'Calibri'
        
        # Save to BytesIO instead of file
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output
